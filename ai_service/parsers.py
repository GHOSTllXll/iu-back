# backend/ai_service/parsers.py
import os
import pdfplumber
import pandas as pd
from docx import Document

# Files larger than this are rejected before parsing (bytes)
MAX_FILE_SIZE = 25 * 1024 * 1024  # 25 MB

ALLOWED_DOC_EXTENSIONS = {'.pdf', '.doc', '.docx'}
ALLOWED_EXCEL_EXTENSIONS = {'.xls', '.xlsx', '.csv'}


class FileValidationError(Exception):
    """Raised when an uploaded file fails basic sanity checks."""
    pass


def validate_uploaded_file(uploaded_file, allowed_extensions, label="File"):
    """
    Basic guardrails every uploaded file should pass before any parser touches it.
    Raises FileValidationError with a clean, user-facing message on failure.
    """
    if uploaded_file is None:
        raise FileValidationError(f"{label} is missing.")

    if uploaded_file.size == 0:
        raise FileValidationError(f"{label} is empty (0 bytes). Please re-upload.")

    if uploaded_file.size > MAX_FILE_SIZE:
        size_mb = uploaded_file.size / (1024 * 1024)
        raise FileValidationError(
            f"{label} is {size_mb:.1f}MB, which exceeds the {MAX_FILE_SIZE // (1024*1024)}MB limit."
        )

    ext = os.path.splitext(uploaded_file.name)[1].lower()
    if ext not in allowed_extensions:
        raise FileValidationError(
            f"{label} has an unsupported extension '{ext}'. "
            f"Allowed: {', '.join(sorted(allowed_extensions))}"
        )

    return ext


def extract_text_from_pdf(pdf_file) -> str:
    """
    Extracts all text and basic table structures from a PDF file.
    """
    text_content = []

    try:
        with pdfplumber.open(pdf_file) as pdf:
            for page_num, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text:
                    text_content.append(f"--- PAGE {page_num + 1} ---\n{text}")

                tables = page.extract_tables()
                for table in tables:
                    if not table or not table[0]:
                        continue
                    df = pd.DataFrame(table[1:], columns=table[0])
                    text_content.append(f"\n[TABLE ON PAGE {page_num + 1}]\n{df.to_string(index=False)}\n")
    except Exception as e:
        raise FileValidationError(f"Could not read PDF '{pdf_file.name}': {str(e)}")

    if not text_content:
        raise FileValidationError(
            f"No extractable text found in '{pdf_file.name}'. "
            "It may be a scanned/image-only PDF that needs OCR."
        )

    return "\n".join(text_content)


def extract_text_from_docx(docx_file) -> str:
    """
    Extracts text and table content from a Word (.docx) file.
    """
    try:
        doc = Document(docx_file)
    except Exception as e:
        raise FileValidationError(f"Could not read Word document '{docx_file.name}': {str(e)}")

    text_content = []

    for para in doc.paragraphs:
        if para.text.strip():
            text_content.append(para.text)

    for table_num, table in enumerate(doc.tables):
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if rows:
            df = pd.DataFrame(rows[1:], columns=rows[0]) if len(rows) > 1 else pd.DataFrame(rows)
            text_content.append(f"\n[TABLE {table_num + 1}]\n{df.to_string(index=False)}\n")

    if not text_content:
        raise FileValidationError(f"No extractable text found in '{docx_file.name}'.")

    return "\n".join(text_content)


def extract_document_text(doc_file) -> str:
    """
    Routes a document (OM or T12) to the right parser based on its extension.
    Supports: .pdf, .doc, .docx
    """
    ext = validate_uploaded_file(doc_file, ALLOWED_DOC_EXTENSIONS, label=doc_file.name)

    if ext == '.pdf':
        return extract_text_from_pdf(doc_file)
    elif ext in ('.doc', '.docx'):
        return extract_text_from_docx(doc_file)

    # Should be unreachable given validate_uploaded_file, but fail loudly if it happens
    raise FileValidationError(f"Unhandled document extension '{ext}' for '{doc_file.name}'.")


def extract_t12_text(t12_file) -> str:
    """
    T12 can legitimately be a PDF/Word doc OR an Excel/CSV file — route based on extension.
    """
    ext = os.path.splitext(t12_file.name)[1].lower()

    if ext in ALLOWED_EXCEL_EXTENSIONS:
        return extract_data_from_excel(t12_file)
    elif ext in ALLOWED_DOC_EXTENSIONS:
        return extract_document_text(t12_file)
    else:
        raise FileValidationError(
            f"'{t12_file.name}' has an unsupported extension '{ext}'. "
            f"Allowed: {', '.join(sorted(ALLOWED_DOC_EXTENSIONS | ALLOWED_EXCEL_EXTENSIONS))}"
        )


def detect_header_row(excel_file, max_scan_rows: int = 15) -> int:
    """
    Real-world spreadsheets often have a title row (e.g. a merged "Rent Roll"
    heading) above the actual column headers, which fools pandas' default
    assumption that row 0 is the header — resulting in a DataFrame full of
    'Unnamed: N' columns and one real column name from the title text.

    Heuristic: scan the first few rows and pick whichever has the most
    non-empty cells. A genuine header row has most/all columns populated with
    short labels; a title row typically has just one or two cells filled in.
    This isn't foolproof, but it correctly handles the common "title row
    above real headers" case.

    Returns the 0-indexed row number to use as the header.
    """
    excel_file.seek(0)
    raw = pd.read_excel(excel_file, sheet_name=0, header=None, nrows=max_scan_rows)
    excel_file.seek(0)

    best_row_idx = 0
    best_non_null_count = -1

    for i in range(len(raw)):
        non_null_count = raw.iloc[i].notna().sum()
        if non_null_count > best_non_null_count:
            best_non_null_count = non_null_count
            best_row_idx = i

    return best_row_idx


def _dedupe_columns(columns):
    """
    Mimics pandas' own behavior for duplicate column names (e.g. two columns
    both literally named 'Unit' become 'Unit' and 'Unit.1') — needed because
    we're building column names manually here rather than letting pandas'
    header= parameter handle it for us.
    """
    seen = {}
    result = []
    for col in columns:
        if col not in seen:
            seen[col] = 0
            result.append(col)
        else:
            seen[col] += 1
            result.append(f"{col}.{seen[col]}")
    return result


def _build_column_names(raw: pd.DataFrame, header_row_idx: int):
    """
    Builds column names from the detected header row, merging in a sub-header
    row directly below it if one looks present. Handles the common
    "grouped header" pattern in real-world rent rolls, e.g. a cell reading
    "Market" with "Rent" directly beneath it in the next row — meaning
    "Market Rent" — which a single-row header read would otherwise split
    apart, losing the "Rent" qualifier entirely (and potentially causing a
    genuine rent column to go undetected downstream).

    Returns (column_names, used_subheader: bool).
    """
    header_row = raw.iloc[header_row_idx]
    next_row_idx = header_row_idx + 1
    has_next_row = next_row_idx < len(raw)
    next_row = raw.iloc[next_row_idx] if has_next_row else None

    # Heuristic: the row below only counts as a sub-header if it has a
    # modest number of populated cells (at least 2, but not the majority of
    # the row) — a real data row is usually either much fuller (many
    # populated fields) or, for a blank/vacant row, much sparser than this.
    subheader_count = next_row.notna().sum() if has_next_row else 0
    looks_like_subheader = has_next_row and 2 <= subheader_count <= max(2, len(next_row) * 0.6)

    columns = []
    for i in range(len(header_row)):
        top_val = header_row.iloc[i]
        top_str = str(top_val).strip() if pd.notna(top_val) else ''

        if looks_like_subheader:
            bottom_val = next_row.iloc[i]
            bottom_str = str(bottom_val).strip() if pd.notna(bottom_val) else ''
            if top_str and bottom_str:
                columns.append(f"{top_str} {bottom_str}")
            elif top_str:
                columns.append(top_str)
            elif bottom_str:
                columns.append(bottom_str)
            else:
                columns.append(f"Unnamed: {i}")
        else:
            columns.append(top_str if top_str else f"Unnamed: {i}")

    return _dedupe_columns(columns), looks_like_subheader


def read_excel_with_header_detection(excel_file) -> pd.DataFrame:
    """
    Reads an Excel file into a DataFrame using a detected header row instead
    of blindly assuming row 0 — and, where present, merges a directly-below
    sub-header row into the column names (see _build_column_names). CSV files
    don't have either problem in practice, so this is Excel-only; callers
    should use pd.read_csv directly for .csv files.
    """
    header_row_idx = detect_header_row(excel_file)

    # Read just enough rows raw (no header) to inspect the header + potential
    # sub-header row and build correct column names.
    excel_file.seek(0)
    header_probe = pd.read_excel(excel_file, sheet_name=0, header=None, nrows=header_row_idx + 2)
    excel_file.seek(0)

    columns, used_subheader = _build_column_names(header_probe, header_row_idx)

    data_start_row = header_row_idx + (2 if used_subheader else 1)
    df = pd.read_excel(excel_file, sheet_name=0, header=None, skiprows=data_start_row)
    excel_file.seek(0)

    # Guard against a width mismatch (shouldn't normally happen, but a
    # malformed sheet with inconsistent row widths could theoretically cause
    # one) rather than crashing with a confusing pandas error.
    if len(df.columns) != len(columns):
        columns = columns[:len(df.columns)] + [f"Unnamed: {i}" for i in range(len(columns), len(df.columns))]

    df.columns = columns
    return df


def extract_data_from_excel(excel_file) -> str:
    """
    Extracts data from an Excel/CSV file and formats it as a clean string.
    Resets the file pointer to 0 before AND after reading, so callers can
    safely re-read the same file object afterward (e.g. into a DataFrame).
    """
    ext = validate_uploaded_file(excel_file, ALLOWED_EXCEL_EXTENSIONS, label=excel_file.name)

    try:
        excel_file.seek(0)
        if ext == '.csv':
            df = pd.read_csv(excel_file)
        else:
            df = read_excel_with_header_detection(excel_file)

        df = df.dropna(how='all').dropna(axis=1, how='all')

        if df.empty:
            raise FileValidationError(f"'{excel_file.name}' has no usable data after cleaning empty rows/columns.")

        result = df.to_string(index=False)
    except FileValidationError:
        raise
    except Exception as e:
        raise FileValidationError(f"Could not read '{excel_file.name}': {str(e)}")
    finally:
        # Leave the pointer reset so the caller can safely read it again
        excel_file.seek(0)

    return result